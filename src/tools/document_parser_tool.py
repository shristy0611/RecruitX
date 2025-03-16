from typing import Any, Dict, Optional
import fitz
from docx import Document
from pathlib import Path
from pydantic import BaseModel, Field

from app.tool.base import BaseTool, ToolResult

class DocumentParserInput(BaseModel):
    """Input schema for document parsing."""
    file_path: str = Field(..., description="Path to the document file")
    extract_metadata: bool = Field(default=True, description="Whether to extract metadata")

class DocumentParserTool(BaseTool):
    """Tool for parsing different document formats using PyMuPDF and python-docx."""
    
    name: str = "document_parser"
    description: str = "Parse documents (PDF, DOCX) and extract text and metadata"
    parameters: Dict = {
        "type": "object",
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Path to the document file"
            },
            "extract_metadata": {
                "type": "boolean",
                "description": "Whether to extract metadata",
                "default": True
            }
        },
        "required": ["file_path"]
    }
    
    async def execute(self, **kwargs) -> ToolResult:
        """Execute document parsing."""
        try:
            input_data = DocumentParserInput(**kwargs)
            file_path = Path(input_data.file_path)
            
            if not file_path.exists():
                return ToolResult(error=f"File not found: {file_path}")
                
            file_type = file_path.suffix.lower()
            
            if file_type == '.pdf':
                result = await self._parse_pdf(file_path, input_data.extract_metadata)
            elif file_type == '.docx':
                result = await self._parse_docx(file_path, input_data.extract_metadata)
            else:
                return ToolResult(error=f"Unsupported file type: {file_type}")
                
            return ToolResult(output=result)
            
        except Exception as e:
            return ToolResult(error=f"Error parsing document: {str(e)}")
    
    async def _parse_pdf(self, file_path: Path, extract_metadata: bool) -> Dict[str, Any]:
        """Parse a PDF document."""
        doc = fitz.open(file_path)
        text = ""
        metadata = {}
        
        try:
            # Extract text
            for page in doc:
                text += page.get_text()
            
            # Extract metadata if requested
            if extract_metadata:
                metadata = doc.metadata
                
            return {
                'text': text,
                'metadata': metadata,
                'file_type': '.pdf',
                'page_count': len(doc)
            }
            
        finally:
            doc.close()
    
    async def _parse_docx(self, file_path: Path, extract_metadata: bool) -> Dict[str, Any]:
        """Parse a DOCX document."""
        doc = Document(file_path)
        text = ""
        metadata = {}
        
        # Extract text
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Extract metadata if requested
        if extract_metadata:
            metadata = {
                'author': doc.core_properties.author,
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
                'title': doc.core_properties.title
            }
            
        return {
            'text': text,
            'metadata': metadata,
            'file_type': '.docx',
            'paragraph_count': len(doc.paragraphs)
        } 